import os
from docx import Document
import openpyxl

def read_docx(file_path):
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                text.append(" | ".join(row_data))
        return "\n".join(text)
    except Exception as e:
        return f"Error reading {file_path}: {e}"

def read_xlsx(file_path):
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        text = []
        for sheet_name in wb.sheetnames:
            text.append(f"Sheet: {sheet_name}")
            ws = wb[sheet_name]
            for row in ws.iter_rows(values_only=True):
                # filter None
                row_data = [str(cell) for cell in row if cell is not None]
                if row_data:
                    text.append(" | ".join(row_data))
        return "\n".join(text)
    except Exception as e:
        return f"Error reading {file_path}: {e}"

files = [
    "AAD - A8 Project Charter.docx",
    "AAD - RACI and Stakeholder Register.xlsx",
    "AAD - RAID Log.xlsx",
    "Hubspot Wishlist (1).docx"
]

with open("content_dump.md", "w", encoding="utf-8") as f:
    for filename in files:
        f.write(f"\n\n# File: {filename}\n")
        f.write("-" * 40 + "\n")
        if filename.endswith(".docx"):
            content = read_docx(filename)
        elif filename.endswith(".xlsx"):
            content = read_xlsx(filename)
        else:
            content = "Unsupported file type"
        f.write(content)
